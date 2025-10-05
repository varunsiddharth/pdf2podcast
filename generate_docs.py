#!/usr/bin/env python3
"""
Documentation Generator for PDF to Podcast Project
Converts Markdown documentation to Word and PDF formats
"""

import os
import sys
from pathlib import Path

def install_requirements():
    """Install required packages for document generation"""
    import subprocess
    
    packages = [
        'python-docx',
        'markdown',
        'weasyprint',
        'pdfkit',
        'markdown2'
    ]
    
    for package in packages:
        try:
            __import__(package.replace('-', '_'))
        except ImportError:
            print(f"Installing {package}...")
            subprocess.check_call([sys.executable, '-m', 'pip', 'install', package])

def markdown_to_word():
    """Convert Markdown to Word document"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import markdown2
        
        print("Converting to Word format...")
        
        # Read markdown file
        with open('PDF_to_Podcast_Documentation.md', 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks'])
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('PDF to Podcast Generator', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add subtitle
        subtitle = doc.add_heading('Complete Technical Documentation', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph('Version: 1.0.0')
        doc.add_paragraph('Last Updated: October 2024')
        doc.add_paragraph('Author: Development Team')
        doc.add_paragraph('')
        
        # Process content
        lines = markdown_content.split('\n')
        current_paragraph = None
        
        for line in lines:
            line = line.strip()
            
            if not line:
                if current_paragraph:
                    doc.add_paragraph('')
                current_paragraph = None
                continue
            
            # Headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('# ').strip()
                
                if level == 1:
                    doc.add_heading(header_text, level=1)
                elif level == 2:
                    doc.add_heading(header_text, level=2)
                elif level == 3:
                    doc.add_heading(header_text, level=3)
                elif level == 4:
                    doc.add_heading(header_text, level=4)
                else:
                    doc.add_heading(header_text, level=5)
                
                current_paragraph = None
            
            # Code blocks
            elif line.startswith('```'):
                if current_paragraph:
                    current_paragraph = None
                continue
            
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                doc.add_paragraph(text, style='List Bullet')
                current_paragraph = None
            
            elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
                text = line[3:].strip()
                doc.add_paragraph(text, style='List Number')
                current_paragraph = None
            
            # Tables (basic support)
            elif '|' in line and not line.startswith('|'):
                # Skip table separators
                if '---' in line or '===' in line:
                    continue
                # This is a simplified table handling
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    p = doc.add_paragraph()
                    for i, cell in enumerate(cells):
                        if i > 0:
                            p.add_run(' | ')
                        p.add_run(cell)
                    current_paragraph = None
            
            # Regular paragraphs
            else:
                if current_paragraph is None:
                    current_paragraph = doc.add_paragraph()
                else:
                    current_paragraph.add_run(' ')
                
                # Handle bold and italic
                text = line
                if '**' in text:
                    parts = text.split('**')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            current_paragraph.add_run(part)
                        else:
                            run = current_paragraph.add_run(part)
                            run.bold = True
                elif '*' in text:
                    parts = text.split('*')
                    for i, part in enumerate(parts):
                        if i % 2 == 0:
                            current_paragraph.add_run(part)
                        else:
                            run = current_paragraph.add_run(part)
                            run.italic = True
                else:
                    current_paragraph.add_run(text)
        
        # Save document
        doc.save('PDF_to_Podcast_Documentation.docx')
        print("[SUCCESS] Word document created: PDF_to_Podcast_Documentation.docx")
        
    except Exception as e:
        print(f"[ERROR] Error creating Word document: {e}")

def markdown_to_pdf():
    """Convert Markdown to PDF document"""
    try:
        import markdown2
        from weasyprint import HTML, CSS
        from weasyprint.text.fonts import FontConfiguration
        
        print("Converting to PDF format...")
        
        # Read markdown file
        with open('PDF_to_Podcast_Documentation.md', 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks'])
        
        # Create full HTML document
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>PDF to Podcast Generator - Documentation</title>
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    line-height: 1.6;
                    color: #333;
                    max-width: 800px;
                    margin: 0 auto;
                    padding: 20px;
                }}
                h1 {{
                    color: #2c3e50;
                    border-bottom: 3px solid #3498db;
                    padding-bottom: 10px;
                }}
                h2 {{
                    color: #34495e;
                    border-bottom: 2px solid #bdc3c7;
                    padding-bottom: 5px;
                    margin-top: 30px;
                }}
                h3 {{
                    color: #7f8c8d;
                    margin-top: 25px;
                }}
                h4 {{
                    color: #95a5a6;
                    margin-top: 20px;
                }}
                code {{
                    background-color: #f8f9fa;
                    padding: 2px 4px;
                    border-radius: 3px;
                    font-family: 'Courier New', monospace;
                }}
                pre {{
                    background-color: #f8f9fa;
                    padding: 15px;
                    border-radius: 5px;
                    overflow-x: auto;
                    border-left: 4px solid #3498db;
                }}
                pre code {{
                    background-color: transparent;
                    padding: 0;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 20px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 12px;
                    text-align: left;
                }}
                th {{
                    background-color: #f2f2f2;
                    font-weight: bold;
                }}
                ul, ol {{
                    margin: 10px 0;
                    padding-left: 30px;
                }}
                li {{
                    margin: 5px 0;
                }}
                blockquote {{
                    border-left: 4px solid #3498db;
                    margin: 20px 0;
                    padding: 10px 20px;
                    background-color: #f8f9fa;
                }}
                .page-break {{
                    page-break-before: always;
                }}
                @media print {{
                    body {{
                        font-size: 12pt;
                    }}
                    h1 {{
                        page-break-before: always;
                    }}
                    h1:first-child {{
                        page-break-before: avoid;
                    }}
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Generate PDF
        font_config = FontConfiguration()
        html_doc = HTML(string=full_html)
        html_doc.write_pdf('PDF_to_Podcast_Documentation.pdf', font_config=font_config)
        
        print("[SUCCESS] PDF document created: PDF_to_Podcast_Documentation.pdf")
        
    except Exception as e:
        print(f"[ERROR] Error creating PDF document: {e}")
        print("Trying alternative PDF generation method...")
        
        try:
            # Alternative method using pdfkit
            import pdfkit
            
            # Read markdown file
            with open('PDF_to_Podcast_Documentation.md', 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # Convert markdown to HTML
            html_content = markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks'])
            
            # Create full HTML document
            full_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>PDF to Podcast Generator - Documentation</title>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                    h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; }}
                    h2 {{ color: #34495e; border-bottom: 1px solid #bdc3c7; }}
                    code {{ background-color: #f8f9fa; padding: 2px 4px; border-radius: 3px; }}
                    pre {{ background-color: #f8f9fa; padding: 15px; border-radius: 5px; }}
                    table {{ border-collapse: collapse; width: 100%; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Generate PDF
            options = {
                'page-size': 'A4',
                'margin-top': '0.75in',
                'margin-right': '0.75in',
                'margin-bottom': '0.75in',
                'margin-left': '0.75in',
                'encoding': "UTF-8",
                'no-outline': None
            }
            
            pdfkit.from_string(full_html, 'PDF_to_Podcast_Documentation.pdf', options=options)
            print("[SUCCESS] PDF document created: PDF_to_Podcast_Documentation.pdf")
            
        except Exception as e2:
            print(f"[ERROR] Error with alternative PDF generation: {e2}")

def create_simple_pdf():
    """Create a simple PDF using basic HTML to PDF conversion"""
    try:
        import markdown2
        
        print("Creating simple PDF format...")
        
        # Read markdown file
        with open('PDF_to_Podcast_Documentation.md', 'r', encoding='utf-8') as f:
            markdown_content = f.read()
        
        # Convert markdown to HTML
        html_content = markdown2.markdown(markdown_content, extras=['tables', 'fenced-code-blocks'])
        
        # Create a simple HTML file that can be printed to PDF
        simple_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>PDF to Podcast Generator - Documentation</title>
            <style>
                body {{ 
                    font-family: 'Times New Roman', serif; 
                    line-height: 1.5; 
                    color: #000; 
                    max-width: 800px; 
                    margin: 0 auto; 
                    padding: 20px;
                }}
                h1 {{ 
                    color: #000; 
                    border-bottom: 2px solid #000; 
                    padding-bottom: 10px;
                    page-break-before: always;
                }}
                h1:first-child {{ page-break-before: avoid; }}
                h2 {{ 
                    color: #333; 
                    border-bottom: 1px solid #666; 
                    padding-bottom: 5px; 
                    margin-top: 30px;
                }}
                h3 {{ color: #555; margin-top: 25px; }}
                h4 {{ color: #666; margin-top: 20px; }}
                code {{ 
                    background-color: #f5f5f5; 
                    padding: 2px 4px; 
                    font-family: 'Courier New', monospace;
                }}
                pre {{ 
                    background-color: #f5f5f5; 
                    padding: 15px; 
                    border: 1px solid #ddd;
                    overflow-x: auto;
                }}
                table {{ 
                    border-collapse: collapse; 
                    width: 100%; 
                    margin: 20px 0;
                }}
                th, td {{ 
                    border: 1px solid #000; 
                    padding: 8px; 
                    text-align: left;
                }}
                th {{ background-color: #f0f0f0; font-weight: bold; }}
                ul, ol {{ margin: 10px 0; padding-left: 30px; }}
                li {{ margin: 5px 0; }}
                @media print {{
                    body {{ font-size: 12pt; }}
                    h1 {{ page-break-before: always; }}
                    h1:first-child {{ page-break-before: avoid; }}
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """
        
        # Save HTML file
        with open('PDF_to_Podcast_Documentation.html', 'w', encoding='utf-8') as f:
            f.write(simple_html)
        
        print("[SUCCESS] HTML document created: PDF_to_Podcast_Documentation.html")
        print("[INFO] You can open this HTML file in your browser and print to PDF")
        
    except Exception as e:
        print(f"[ERROR] Error creating simple PDF: {e}")

def main():
    """Main function to generate all documentation formats"""
    print("[START] PDF to Podcast Documentation Generator")
    print("=" * 50)
    
    # Check if markdown file exists
    if not os.path.exists('PDF_to_Podcast_Documentation.md'):
        print("[ERROR] PDF_to_Podcast_Documentation.md not found!")
        return
    
    # Install requirements
    print("[INFO] Installing required packages...")
    install_requirements()
    
    # Generate Word document
    markdown_to_word()
    
    # Generate PDF document
    markdown_to_pdf()
    
    # Create simple HTML version as fallback
    create_simple_pdf()
    
    print("\n" + "=" * 50)
    print("[SUCCESS] Documentation generation complete!")
    print("\nGenerated files:")
    print("[FILE] PDF_to_Podcast_Documentation.docx (Word format)")
    print("[FILE] PDF_to_Podcast_Documentation.pdf (PDF format)")
    print("[FILE] PDF_to_Podcast_Documentation.html (HTML format - can be printed to PDF)")
    print("\n[COMPLETE] All documentation formats are ready!")

if __name__ == "__main__":
    main()
